from pathlib import Path
import re
import unicodedata
from typing import Optional

import openpyxl
import pandas as pd
from docx import Document
from openpyxl.cell.cell import MergedCell


base_dir = Path(__file__).resolve().parent
data_dir = base_dir / "data"
exit_dir = base_dir / "exit"
temp_dir = base_dir / "temp_inputs"

exit_dir.mkdir(exist_ok=True)
temp_dir.mkdir(exist_ok=True)

arquivo_consemavi_padrao = data_dir / "PDM 2025-2028 - Meta 49 - Recapeamento - Janeiro.26.xlsx"
arquivo_modelo_padrao = data_dir / "Meta49-formatado.xlsx"
arquivo_word_padrao = data_dir / "Meta 49 - Janeiro-2026 - CONVIAS.docx"
arquivo_saida_padrao = exit_dir / "meta49_preenchido.xlsx"

subpref_map = {
    "ad": "cidade ademar",
    "af": "aricanduva/formosa/carrão",
    "bt": "butantã",
    "cl": "campo limpo",
    "cs": "capela do socorro",
    "ct": "cidade tiradentes",
    "cv": "casa verde/cachoeirinha",
    "em": "ermelino matarazzo",
    "fb": "freguesia do ó/brasilândia",
    "g": "guaianases",
    "ip": "ipiranga",
    "iq": "itaquera",
    "it": "itaim paulista",
    "ja": "jabaquara",
    "jt": "jaçanã/tremembé",
    "la": "lapa",
    "mb": "m'boi mirim",
    "mg": "vila maria/vila guilherme",
    "mo": "mooca",
    "mp": "são miguel paulista",
    "pa": "parelheiros",
    "pe": "penha",
    "pi": "pinheiros",
    "pj": "pirituba/jaraguá",
    "pr": "perus/anhanguera",
    "sa": "santo amaro",
    "sb": "sapopemba",
    "se": "sé",
    "sm": "são mateus",
    "st": "santana/tucuruvi",
    "vm": "vila mariana",
    "vp": "vila prudente",
    "dzu": "departamento de zeladoria urbana",
}

meses_map = {
    "janeiro": "jan",
    "fevereiro": "fev",
    "março": "mar",
    "abril": "abr",
    "maio": "mai",
    "junho": "jun",
    "julho": "jul",
    "agosto": "ago",
    "setembro": "set",
    "outubro": "out",
    "novembro": "nov",
    "dezembro": "dez",
}
meses = list(meses_map.keys())

sigla_aliases = {
    "fb": "fo",
}


def padronizar_texto(texto) -> str:
    if pd.isna(texto):
        return ""
    texto = str(texto).strip().lower()
    texto = unicodedata.normalize("NFKD", texto).encode("ASCII", "ignore").decode("ASCII")
    texto = re.sub(r"\s+", " ", texto)
    return texto


def normalizar_sigla(sigla: str) -> str:
    if pd.isna(sigla):
        return ""

    s = str(sigla).strip().lower()
    s = re.sub(r"\s+", "", s)

    return sigla_aliases.get(s, s)


def traduzir_sub(sigla: str) -> str:
    return subpref_map.get(normalizar_sigla(sigla), sigla)


def numero_br_para_float(valor) -> float:
    if pd.isna(valor):
        return 0.0

    # se já é número de verdade, retorna direto
    if isinstance(valor, (int, float)):
        return float(valor)

    texto = str(valor).strip()

    if texto in {"", "-", "–", "—", "nan", "None"}:
        return 0.0

    texto = texto.replace("m²", "").replace("M²", "").replace("m2", "").strip()

    # caso 1: formato brasileiro -> 3.041,30
    if "," in texto:
        texto = texto.replace(".", "").replace(",", ".")

    # caso 2: formato normal -> 3041.30
    texto = re.sub(r"[^0-9.\-]", "", texto)

    if texto in {"", "-", ".", "-."}:
        return 0.0

    try:
        return float(texto)
    except ValueError:
        return 0.0

def safe_float(valor) -> float:
    try:
        return float(valor)
    except Exception:
        return 0.0


def classificar_indicador(texto: str) -> Optional[str]:
    t = padronizar_texto(texto)

    if "convias" in t:
        return "convias"
    if "consemavi" in t:
        return "consemavi"
    if "total" in t and "requalificado" in t:
        return "total"

    return None


def extrair_textos_docx(caminho_word: Path) -> list[str]:
    doc = Document(caminho_word)
    textos = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            textos.append(t)

    for tabela in doc.tables:
        for row in tabela.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    textos.append(t)

    return textos


def ler_word_convias(caminho_word: Path, mes_ref: str) -> pd.DataFrame:
    textos = extrair_textos_docx(caminho_word)
    textos_limpos = [t.strip() for t in textos if str(t).strip()]

    dados = []
    i = 0

    while i < len(textos_limpos):
        atual = padronizar_texto(textos_limpos[i])

        if atual == "total":
            break

        sigla_valida = re.fullmatch(r"[a-z]{1,3}", atual) is not None

        if sigla_valida:
            sigla = normalizar_sigla(atual)

            if i + 1 < len(textos_limpos):
                proximo = textos_limpos[i + 1].strip()
                valor = numero_br_para_float(proximo)

                dados.append({
                    "sigla": sigla,
                    "mês": mes_ref,
                    "convias": valor
                })
                i += 2
                continue

        i += 1

    df = pd.DataFrame(dados)

    if df.empty:
        raise ValueError("nenhum dado do convias foi lido do word. verifique se o docx está no formato esperado.")

    df["sigla"] = df["sigla"].apply(normalizar_sigla)
    df["convias"] = df["convias"].apply(numero_br_para_float)
    df = df.groupby(["sigla", "mês"], as_index=False)["convias"].sum()

    return df


def montar_df_consemavi(ano_ref: int, caminho_excel: Path) -> pd.DataFrame:
    df_raw = pd.read_excel(caminho_excel, sheet_name="49", header=None)

    # pega a coluna onde normalmente ficam os títulos
    col_titulo = df_raw.iloc[:, 2].astype(str).apply(padronizar_texto)

    titulo_ano = f"area recapeada em {ano_ref}"

    # localizar início do bloco do ano
    linhas_inicio = col_titulo[col_titulo == titulo_ano].index.tolist()
    if not linhas_inicio:
        raise ValueError(f"título 'área recapeada em {ano_ref}' não encontrado no excel do consemavi.")

    linha_titulo_idx = linhas_inicio[0]

    # localizar cabeçalho
    linha_cabecalho_idx = None

    for idx in range(linha_titulo_idx, min(linha_titulo_idx + 5, len(df_raw))):
        row_values = [padronizar_texto(x) for x in df_raw.iloc[idx].tolist()]
        tem_sub = any(x == "sub" for x in row_values)
        tem_janeiro = any(x == "janeiro" for x in row_values)

        if tem_sub and tem_janeiro:
            linha_cabecalho_idx = idx
            break

    if linha_cabecalho_idx is None:
        for idx in range(max(0, linha_titulo_idx - 5), linha_titulo_idx):
            row_values = [padronizar_texto(x) for x in df_raw.iloc[idx].tolist()]
            tem_sub = any(x == "sub" for x in row_values)
            tem_janeiro = any(x == "janeiro" for x in row_values)

            if tem_sub and tem_janeiro:
                linha_cabecalho_idx = idx
                break

    if linha_cabecalho_idx is None:
        raise ValueError("não foi possível localizar a linha de cabeçalho do consemavi.")

    # localizar fim do bloco do ano
    linhas_fim = col_titulo[col_titulo.str.contains("total", na=False)].index.tolist()

    if not linhas_fim:
        raise ValueError(f"linha de total do ano {ano_ref} não encontrada no excel do consemavi.")

    # pega a primeira linha de total depois do título do ano
    linhas_fim_validas = [i for i in linhas_fim if i > linha_titulo_idx]

    if not linhas_fim_validas:
        raise ValueError(f"não foi encontrada uma linha de total após o bloco do ano {ano_ref}.")

    linha_fim_idx = linhas_fim_validas[0]

    # aplicar cabeçalho e cortar só o bloco certo
    df_bloco = df_raw.iloc[linha_cabecalho_idx + 1:linha_fim_idx].copy()
    df_bloco.columns = df_raw.iloc[linha_cabecalho_idx]

    # localizar coluna da sigla/sub
    col_sigla = None
    for c in df_bloco.columns:
        if "sub" in padronizar_texto(c):
            col_sigla = c
            break

    if col_sigla is None:
        raise ValueError("coluna de subprefeitura/sigla não encontrada no consemavi.")

    # localizar colunas de meses
    colunas_meses = [c for c in df_bloco.columns if padronizar_texto(str(c).strip()) in meses]

    if not colunas_meses:
        raise ValueError("nenhuma coluna de mês foi encontrada no consemavi.")

    df_bloco = df_bloco[[col_sigla] + colunas_meses].copy()
    df_bloco = df_bloco.rename(columns={col_sigla: "sigla"})

    # limpar siglas
    df_bloco["sigla"] = df_bloco["sigla"].astype(str).str.strip()
    df_bloco = df_bloco[df_bloco["sigla"] != ""]
    df_bloco["sigla"] = df_bloco["sigla"].apply(normalizar_sigla)
    df_bloco = df_bloco[df_bloco["sigla"].str.match(r"^[a-z]{1,3}$", na=False)]

    # converter meses
    for mes in colunas_meses:
        df_bloco[mes] = df_bloco[mes].apply(numero_br_para_float)

    # agrupar por sigla
    df_agrupado = df_bloco.groupby("sigla", as_index=False)[colunas_meses].sum()

    # padronizar nomes dos meses
    rename_map = {col: padronizar_texto(col) for col in colunas_meses}
    df_agrupado = df_agrupado.rename(columns=rename_map)
    colunas_meses_norm = [rename_map[col] for col in colunas_meses]

    # formato longo
    df_longo = df_agrupado.melt(
        id_vars=["sigla"],
        value_vars=colunas_meses_norm,
        var_name="mês",
        value_name="consemavi"
    )

    df_longo["consemavi"] = df_longo["consemavi"].apply(numero_br_para_float)

    return df_longo

def montar_df_final(
    caminho_word: Path,
    caminho_consemavi: Path,
    ano_ref: int,
    mes_ref: str
) -> pd.DataFrame:
    df_convias = ler_word_convias(caminho_word, mes_ref)
    df_consemavi = montar_df_consemavi(ano_ref, caminho_consemavi)

    mes_ref_norm = padronizar_texto(mes_ref)

    if mes_ref_norm != "todos os meses":
        df_convias = df_convias[df_convias["mês"] == mes_ref].copy()
        df_consemavi = df_consemavi[df_consemavi["mês"] == mes_ref_norm].copy()

    df_convias["sigla"] = df_convias["sigla"].apply(normalizar_sigla)
    df_consemavi["sigla"] = df_consemavi["sigla"].apply(normalizar_sigla)

    df_final = pd.merge(
        df_convias,
        df_consemavi,
        on=["sigla", "mês"],
        how="outer"
    )

    df_final["convias"] = df_final["convias"].fillna(0).apply(numero_br_para_float)
    df_final["consemavi"] = df_final["consemavi"].fillna(0).apply(numero_br_para_float)
    df_final["total requalificado"] = df_final["convias"] + df_final["consemavi"]
    df_final["subprefeitura"] = df_final["sigla"].apply(traduzir_sub)

    df_final = df_final.sort_values(["sigla", "mês"]).reset_index(drop=True)

    total_geral = pd.DataFrame({
        "subprefeitura": ["total geral"],
        "sigla": ["total"],
        "mês": ["anual" if mes_ref_norm == "todos os meses" else mes_ref],
        "convias": [df_final["convias"].sum()],
        "consemavi": [df_final["consemavi"].sum()],
        "total requalificado": [df_final["total requalificado"].sum()],
    })

    return pd.concat([df_final, total_geral], ignore_index=True)


def localizar_aba_meta49(wb):
    for nome in wb.sheetnames:
        if "49" in str(nome):
            return wb[nome]
    raise ValueError("aba meta 49 não encontrada no arquivo modelo.")


def localizar_colunas(ws):
    mapa_meses = {}
    col_total = None

    for row in ws.iter_rows(min_row=1, max_row=12):
        for cell in row:
            val = padronizar_texto(cell.value)

            for nome_completo, abrev in meses_map.items():
                if val == padronizar_texto(abrev):
                    mapa_meses[padronizar_texto(nome_completo)] = cell.column

            if val == "total":
                col_total = cell.column

    if not mapa_meses:
        raise ValueError("não foi possível localizar as colunas de meses no modelo.")
    if col_total is None:
        raise ValueError("não foi possível localizar a coluna total no modelo.")

    return mapa_meses, col_total


def localizar_linhas(ws):
    mapa = {}
    sigla_atual = None

    for r in range(1, ws.max_row + 1):
        v_sigla = ws.cell(r, 2).value
        v_ind = ws.cell(r, 3).value

        if v_sigla:
            sigla_atual = normalizar_sigla(v_sigla)

        ind_tipo = classificar_indicador(v_ind)

        if sigla_atual and ind_tipo:
            mapa[(sigla_atual, ind_tipo)] = r

    return mapa


def localizar_linha_total_meses(ws):
    for r in range(ws.max_row, 1, -1):
        valores = [padronizar_texto(ws.cell(r, c).value) for c in range(1, min(ws.max_column, 6) + 1)]
        if any("total meses" in v for v in valores):
            return r

    for r in range(ws.max_row, 1, -1):
        valores = [padronizar_texto(ws.cell(r, c).value) for c in range(1, min(ws.max_column, 6) + 1)]
        if any(v == "total" for v in valores):
            return r

    return None


def celula_e_mesclada(ws, linha, coluna):
    return isinstance(ws.cell(linha, coluna), MergedCell)


def limpar_linha(ws, linha):
    if linha is None or linha < 1:
        return

    for c in range(1, ws.max_column + 1):
        if not celula_e_mesclada(ws, linha, c):
            ws.cell(linha, c).value = None


def escrever_se_possivel(ws, linha, coluna, valor):
    if linha is None or coluna is None:
        return
    if not celula_e_mesclada(ws, linha, coluna):
        ws.cell(linha, coluna).value = valor


def aplicar_formato_brasileiro(ws, linha, coluna):
    if linha is None or coluna is None:
        return
    if not celula_e_mesclada(ws, linha, coluna):
        ws.cell(linha, coluna).number_format = '#,##0.00'

def preencher_excel_formatado(
    df_final: pd.DataFrame,
    caminho_modelo: Path,
    caminho_saida: Path
) -> Path:
    # 1. Carrega o arquivo mantendo a formatação original
    wb = openpyxl.load_workbook(caminho_modelo)
    ws = localizar_aba_meta49(wb)

    # 2. Mapeia onde as colunas (meses) e linhas (siglas/indicadores) estão
    mapa_meses, col_total = localizar_colunas(ws)
    mapa_linhas = localizar_linhas(ws)

    # 3. Preenche os dados do DataFrame no Excel
    for _, row in df_final.iterrows():
        sigla = normalizar_sigla(row["sigla"])
        mes = padronizar_texto(row["mês"])

        # Ignora a linha de total do DataFrame, pois vamos recalcular no Excel
        if sigla == "total":
            continue

        col = mapa_meses.get(mes)
        if not col:
            continue

        # Distribui os valores nas linhas correspondentes (Total, Convias, Consemavi)
        mapeamento_valores = [
            ("total", "total requalificado"),
            ("convias", "convias"),
            ("consemavi", "consemavi"),
        ]

        for tipo, coluna_origem in mapeamento_valores:
            linha_excel = mapa_linhas.get((sigla, tipo))
            if linha_excel:
                escrever_se_possivel(ws, linha_excel, col, row[coluna_origem])

    # 4. Recalcula os totais horizontais (Soma dos meses por linha)
    for (sigla, tipo), lin in mapa_linhas.items():
        if sigla == "total":
            continue

        soma_linha = sum(safe_float(ws.cell(lin, c).value) for c in mapa_meses.values())
        escrever_se_possivel(ws, lin, col_total, soma_linha)

    # 5. Localiza a linha de rodapé "Total Meses" para somas verticais
    linha_total_meses = localizar_linha_total_meses(ws)

    if linha_total_meses:
        # Limpa possíveis textos residuais de 'total' abaixo da linha principal
        if linha_total_meses + 1 <= ws.max_row:
            limpar_linha(ws, linha_total_meses + 1)

        escrever_se_possivel(ws, linha_total_meses, 1, "TOTAL MESES:")

        soma_geral_anual = 0
        for mes_nome, col_idx in mapa_meses.items():
            total_do_mes = 0
            # Soma apenas as linhas do tipo "total" de cada subprefeitura
            for (sigla, tipo), lin in mapa_linhas.items():
                if sigla != "total" and tipo == "total":
                    total_do_mes += safe_float(ws.cell(lin, col_idx).value)

            escrever_se_possivel(ws, linha_total_meses, col_idx, total_do_mes)
            soma_geral_anual += total_do_mes

        # Preenche o total do total (canto inferior direito)
        escrever_se_possivel(ws, linha_total_meses, col_total, soma_geral_anual)

    # 6. Salva o resultado final
    wb.save(caminho_saida)
    return caminho_saida


def gerar_relatorio_final(
    ano_ref: int = 2025,
    mes_ref: str = "janeiro",
    caminho_word: Optional[Path] = None,
    caminho_consemavi: Optional[Path] = None,
    caminho_saida: Optional[Path] = None,
) -> Path:
    caminho_word = Path(caminho_word) if caminho_word else arquivo_word_padrao
    caminho_consemavi = Path(caminho_consemavi) if caminho_consemavi else arquivo_consemavi_padrao
    caminho_modelo = arquivo_modelo_padrao
    caminho_saida = Path(caminho_saida) if caminho_saida else arquivo_saida_padrao

    if not caminho_word.exists():
        raise FileNotFoundError(f"word do convias não encontrado: {caminho_word}")
    if not caminho_consemavi.exists():
        raise FileNotFoundError(f"excel do consemavi não encontrado: {caminho_consemavi}")
    if not caminho_modelo.exists():
        raise FileNotFoundError(f"arquivo modelo não encontrado: {caminho_modelo}")

    df_final = montar_df_final(
        caminho_word=caminho_word,
        caminho_consemavi=caminho_consemavi,
        ano_ref=ano_ref,
        mes_ref=mes_ref,
    )

    return preencher_excel_formatado(
        df_final=df_final,
        caminho_modelo=caminho_modelo,
        caminho_saida=caminho_saida,
    )


if __name__ == "__main__":
    saida = gerar_relatorio_final(ano_ref=2025, mes_ref="janeiro")
    print(f"gerado em: {saida}")
